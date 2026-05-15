"""
浙江宣传写作客户端 - Flask应用
"""

import os
import re
import io
import sqlite3
from datetime import datetime
from flask import Flask, render_template, request, jsonify, send_file, session, g
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "zjxc-writer-secret-key")

# 默认配置
DEFAULT_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
DEFAULT_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")

# 数据库路径
DATABASE = os.path.join(os.path.dirname(__file__), "articles.db")


def get_db():
    """获取数据库连接"""
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn


@app.teardown_appcontext
def close_db(exception):
    """关闭数据库连接"""
    pass  # 单用户模式下不关闭连接


def init_db():
    """初始化数据库"""
    conn = sqlite3.connect(DATABASE)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT DEFAULT '新对话',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS articles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER DEFAULT 1,
            content TEXT NOT NULL,
            article_type TEXT DEFAULT 'generated',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id INTEGER DEFAULT 1,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    # 创建默认对话
    cursor = conn.execute("SELECT COUNT(*) FROM conversations")
    if cursor.fetchone()[0] == 0:
        conn.execute("INSERT INTO conversations (title) VALUES (?)", ("新对话",))
    conn.commit()
    conn.close()


# 加载系统提示词
with open("prompt/system_prompt.md", "r", encoding="utf-8") as f:
    SYSTEM_PROMPT = f.read()


def get_client(api_key=None):
    """获取API客户端"""
    key = api_key or session.get("api_key") or os.getenv("DEEPSEEK_API_KEY", "")
    base_url = session.get("base_url") or DEFAULT_BASE_URL
    return OpenAI(api_key=key, base_url=base_url)


def get_model():
    """获取模型名称"""
    return session.get("model") or DEFAULT_MODEL


def extract_paragraph_ref(text):
    """从用户输入中提取段落引用"""
    patterns = [
        r"第(\d+)段",
        r"第(\d+)节",
        r"第三段",
        r"第二段",
        r"第一段",
        r"第三部分",
        r"第二部分",
        r"第一部分",
        r"最后一段",
        r"倒数第一段",
        r"最后一部分",
        r"倒数第一部分",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)
    return None


def article_to_docx(content, title="浙江宣传文章"):
    """将文章内容转换为Word文档"""
    doc = Document()

    lines = content.split("\n")
    first_line = lines[0].strip()

    if first_line.startswith("#"):
        first_line = re.sub(r"^#+\s*", "", first_line)
    if first_line.startswith("浙江宣传"):
        first_line = first_line.replace("浙江宣传", "").strip("| ").strip()

    title_para = doc.add_heading(first_line, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    time_para = doc.add_paragraph(
        f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}"
    )
    time_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()

    in_code_block = False
    for line in lines[1:]:
        if line.strip().startswith("链接:") or line.strip().startswith("---"):
            continue
        if not line.strip():
            continue

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = "Quote"
            continue

        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=3)
        else:
            clean_line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            clean_line = re.sub(r"\*(.+?)\*", r"\1", clean_line)
            clean_line = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", clean_line)
            if clean_line.strip():
                doc.add_paragraph(clean_line.strip())

    return doc


@app.route("/")
def index():
    """主页"""
    if "api_key" not in session:
        session["api_key"] = os.getenv("DEEPSEEK_API_KEY", "")
    if "base_url" not in session:
        session["base_url"] = DEFAULT_BASE_URL
    if "model" not in session:
        session["model"] = DEFAULT_MODEL
    if "conversation_id" not in session:
        session["conversation_id"] = 1
    return render_template("index.html")


@app.route("/config", methods=["GET", "POST"])
def config():
    """配置API"""
    if request.method == "POST":
        data = request.json
        api_key = data.get("api_key", "").strip()
        base_url = data.get("base_url", DEFAULT_BASE_URL).strip()
        model = data.get("model", DEFAULT_MODEL).strip()

        if not api_key:
            return jsonify({"error": "API Key不能为空"}), 400

        session["api_key"] = api_key
        session["base_url"] = base_url or DEFAULT_BASE_URL
        session["model"] = model or DEFAULT_MODEL

        # 持久化到.env文件
        env_path = os.path.join(os.path.dirname(__file__), ".env")
        try:
            with open(env_path, "w", encoding="utf-8") as f:
                f.write(f"# DeepSeek API 配置\n")
                f.write(f"DEEPSEEK_API_KEY={api_key}\n")
                f.write(f"DEEPSEEK_BASE_URL={base_url or DEFAULT_BASE_URL}\n")
                f.write(f"DEEPSEEK_MODEL={model or DEFAULT_MODEL}\n")
        except Exception as e:
            pass

        return jsonify({"success": True, "message": "配置已保存"})

    return jsonify(
        {
            "api_key": session.get("api_key", ""),
            "base_url": session.get("base_url", DEFAULT_BASE_URL),
            "model": session.get("model", DEFAULT_MODEL),
            "has_config": bool(session.get("api_key")),
        }
    )


@app.route("/conversations", methods=["GET"])
def get_conversations():
    """获取对话列表"""
    db = get_db()
    rows = db.execute(
        "SELECT id, title, created_at, updated_at FROM conversations ORDER BY updated_at DESC"
    ).fetchall()
    return jsonify(
        {
            "conversations": [
                {
                    "id": row["id"],
                    "title": row["title"],
                    "created_at": row["created_at"],
                    "updated_at": row["updated_at"],
                }
                for row in rows
            ],
            "current_id": session.get("conversation_id", 1),
        }
    )


@app.route("/conversations", methods=["POST"])
def create_conversation():
    """创建新对话"""
    db = get_db()
    cursor = db.execute("INSERT INTO conversations (title) VALUES (?)", ("新对话",))
    new_id = cursor.lastrowid
    db.commit()

    # 清空当前对话的内容
    db.execute("DELETE FROM messages WHERE conversation_id = ?", (new_id,))
    db.execute("DELETE FROM articles WHERE conversation_id = ?", (new_id,))
    db.commit()

    session["conversation_id"] = new_id
    return jsonify({"success": True, "id": new_id})


@app.route("/conversations/<int:conv_id>", methods=["POST"])
def switch_conversation(conv_id):
    """切换对话"""
    db = get_db()
    row = db.execute("SELECT id FROM conversations WHERE id = ?", (conv_id,)).fetchone()
    if not row:
        return jsonify({"error": "对话不存在"}), 404

    session["conversation_id"] = conv_id
    return jsonify({"success": True})


@app.route("/conversations/<int:conv_id>", methods=["DELETE"])
def delete_conversation(conv_id):
    """删除对话"""
    db = get_db()
    if conv_id == 1:
        return jsonify({"error": "默认对话不能删除"}), 400

    db.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
    db.execute("DELETE FROM articles WHERE conversation_id = ?", (conv_id,))
    db.execute("DELETE FROM conversations WHERE id = ?", (conv_id,))
    db.commit()

    if session.get("conversation_id") == conv_id:
        session["conversation_id"] = 1

    return jsonify({"success": True})


@app.route("/conversations/<int:conv_id>/messages", methods=["GET"])
def get_conversation_messages(conv_id):
    """获取指定对话的消息历史"""
    db = get_db()
    rows = db.execute(
        "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY id",
        (conv_id,),
    ).fetchall()
    articles = db.execute(
        "SELECT id, content, article_type, created_at FROM articles WHERE conversation_id = ? ORDER BY id",
        (conv_id,),
    ).fetchall()
    return jsonify(
        {
            "messages": [
                {"role": row["role"], "content": row["content"]} for row in rows
            ],
            "articles": [
                {
                    "id": row["id"],
                    "content": row["content"],
                    "type": row["article_type"],
                    "created_at": row["created_at"],
                }
                for row in articles
            ],
        }
    )


@app.route("/generate", methods=["POST"])
def generate():
    """生成文章"""
    if not session.get("api_key"):
        return jsonify({"error": "请先配置API Key", "need_config": True}), 400

    data = request.json
    user_input = data.get("input", "").strip()

    if not user_input:
        return jsonify({"error": "请输入内容"}), 400

    db = get_db()

    # 模板请求
    if "模板" in user_input:
        topics = """当前可用主题模板：

**经济类**：消费升级、民营经济、制造业转型、平台经济、年轻人创业
**社会类**：延迟退休、生育率下降、县城崛起、就业难、人口流动
**文化类**：国潮崛起、传统文化复兴、短视频与注意力、阅读与学习
**科技类**：人工智能影响、新能源发展、芯片自主、数字经济
**民生类**：房价走势、医疗改革、教育公平、养老问题"""
        return jsonify({"type": "template", "content": topics})

    # 配置查询
    if any(kw in user_input for kw in ["配置", "设置", "模型", "当前"]):
        config_info = f"""当前配置：
- Base URL: {session.get('base_url', DEFAULT_BASE_URL)}
- Model: {session.get('model', DEFAULT_MODEL)}
- API Key: {session.get('api_key', '')[:8] if session.get('api_key') else ''}***"""
        return jsonify({"type": "config", "content": config_info})

    # 清空
    if any(kw in user_input for kw in ["清空", "新对话", "重新开始"]):
        conv_id = session.get("conversation_id", 1)
        db.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
        db.execute("DELETE FROM articles WHERE conversation_id = ?", (conv_id,))
        db.commit()
        return jsonify({"type": "clear", "content": "已清空对话，可开始新的话题。"})

    # 历史查询
    if any(kw in user_input for kw in ["历史", "查看文章"]):
        rows = db.execute(
            "SELECT id, content, article_type, created_at FROM articles ORDER BY id"
        ).fetchall()
        if not rows:
            return jsonify(
                {
                    "type": "history",
                    "content": "还没有生成任何文章，请先输入主题生成文章。",
                }
            )

        history_text = f"已生成 {len(rows)} 篇文章：\n\n"
        for i, row in enumerate(rows, 1):
            title_match = re.search(r"[#*]*(.+?)[#*\n]", row["content"])
            title = title_match.group(1).strip()[:40] if title_match else f"第{i}篇"
            preview = (
                row["content"][:80].replace("\n", " ").replace("#", "").replace("*", "")
            )
            history_text += f"{i}. {title}\n{preview}...\n\n"

        return jsonify({"type": "history", "content": history_text})

    # 构建消息列表：系统提示词 + 对话历史 + 用户输入
    # 让AI自己根据上下文判断是生成新文章还是修改旧文章
    messages = [{"role": "system", "content": SYSTEM_PROMPT}]

    # 从数据库加载对话历史
    conv_id = session.get("conversation_id", 1)
    msg_rows = db.execute(
        "SELECT role, content FROM messages WHERE conversation_id = ? ORDER BY id",
        (conv_id,),
    ).fetchall()
    for row in msg_rows:
        messages.append({"role": row["role"], "content": row["content"]})

    # 直接把用户输入作为新消息，让AI自己判断
    user_msg = user_input

    messages.append({"role": "user", "content": user_msg})

    # 调用API
    client = get_client()
    model = get_model()

    try:
        response = client.chat.completions.create(
            model=model, messages=messages, stream=True
        )

        full_content = ""
        for chunk in response:
            if chunk.choices[0].delta.content:
                full_content += chunk.choices[0].delta.content
    except Exception as e:
        return jsonify({"error": f"API调用失败: {str(e)}"}), 500

    # 从文章内容中提取标题
    title_match = re.search(r"[#*]*(.+?)[#*\n]", full_content)
    title = title_match.group(1).strip() if title_match else f"文章"

    # 保存对话历史（用户输入 + 仅非文章的系统回复）
    db.execute(
        "INSERT INTO messages (conversation_id, role, content) VALUES (?, ?, ?)",
        (conv_id, "user", user_msg),
    )
    # 文章内容不存入 messages 表，避免重复显示
    # 只保存简短的确认消息
    db.execute(
        "INSERT INTO messages (conversation_id, role, content) VALUES (?, ?, ?)",
        (conv_id, "assistant", f"【已生成文章】{title}"),
    )

    # 更新对话标题为第一句用户输入
    conv_row = db.execute(
        "SELECT title FROM conversations WHERE id = ?", (conv_id,)
    ).fetchone()
    if conv_row and conv_row["title"] == "新对话":
        # 截取用户输入的前30字符作为标题
        title = user_input[:30] + "..." if len(user_input) > 30 else user_input
        db.execute("UPDATE conversations SET title = ? WHERE id = ?", (title, conv_id))

    # 保存文章
    if paragraph_ref:
        # 修改模式：更新最后一篇
        last_id = db.execute(
            "SELECT id FROM articles WHERE conversation_id = ? ORDER BY id DESC LIMIT 1",
            (conv_id,),
        ).fetchone()
        if last_id:
            db.execute(
                "UPDATE articles SET content = ?, article_type = ? WHERE id = ?",
                (full_content, "modified", last_id["id"]),
            )
            article_index = last_id["id"]
        else:
            db.execute(
                "INSERT INTO articles (conversation_id, content, article_type) VALUES (?, ?, ?)",
                (conv_id, full_content, "modified"),
            )
            article_index = db.execute("SELECT last_insert_rowid()").fetchone()[0]
    else:
        # 新文章模式
        db.execute(
            "INSERT INTO articles (conversation_id, content, article_type) VALUES (?, ?, ?)",
            (conv_id, full_content, "generated"),
        )
        article_index = db.execute("SELECT last_insert_rowid()").fetchone()[0]

    db.commit()

    # 更新对话的updated_at
    db.execute(
        "UPDATE conversations SET updated_at = CURRENT_TIMESTAMP WHERE id = ?",
        (conv_id,),
    )
    db.commit()

    # 生成Word
    doc = article_to_docx(full_content, title)
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    return jsonify(
        {
            "type": "article",
            "content": full_content,
            "title": title,
            "word_data": doc_bytes.getvalue().hex(),
            "index": article_index,
            "is_modified": paragraph_ref is not None,
        }
    )


@app.route("/download/<int:index>")
def download(index):
    """下载Word文件"""
    db = get_db()
    row = db.execute("SELECT content FROM articles WHERE id = ?", (index,)).fetchone()

    if not row:
        return f"文章不存在 (index={index})", 404

    title_match = re.search(r"[#*]*(.+?)[#*\n]", row["content"])
    title = title_match.group(1).strip() if title_match else f"文章{index}"

    doc = article_to_docx(row["content"], title)
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    filename = f"{title[:20]}_{datetime.now().strftime('%Y%m%d%H%M')}.docx"

    return send_file(
        doc_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )


@app.route("/history")
def history():
    """获取历史"""
    db = get_db()
    conv_id = session.get("conversation_id", 1)
    rows = db.execute(
        "SELECT id, content, created_at FROM articles WHERE conversation_id = ? ORDER BY id",
        (conv_id,),
    ).fetchall()
    return jsonify(
        {
            "history": [
                {
                    "index": row["id"],
                    "title": (
                        re.search(r"[#*]*(.+?)[#*\n]", row["content"])
                        .group(1)
                        .strip()[:40]
                        if re.search(r"[#*]*(.+?)[#*\n]", row["content"])
                        else f"第{row['id']}篇"
                    ),
                    "preview": row["content"][:50]
                    .replace("\n", " ")
                    .replace("#", "")
                    .replace("*", ""),
                }
                for row in rows
            ]
        }
    )


@app.route("/clear", methods=["POST"])
def clear():
    """清空对话"""
    db = get_db()
    conv_id = session.get("conversation_id", 1)
    db.execute("DELETE FROM messages WHERE conversation_id = ?", (conv_id,))
    db.execute("DELETE FROM articles WHERE conversation_id = ?", (conv_id,))
    db.commit()
    return jsonify({"success": True})


# 初始化数据库
init_db()


if __name__ == "__main__":
    app.run(debug=True, port=5000)
